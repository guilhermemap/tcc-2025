'''generate_docx_files.py

Esse script lê arquivos csv de perguntas e respostas de uma pesquisa
e separa em um arquivo .docx para cada respondente

Este programa é parte do trabalho de conclusão de curso do curso de
Bacharelado em Ciência da Computação, na Universidade Federal do Paraná,
dos alunos:

André Luis da Silva Machado - GRR20141403
Guilherme Bettu - GRR

=======================================================================
This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with this program.  If not, see <https://www.gnu.org/licenses/>5.
=======================================================================

Curitiba, Brazil - 2025
'''
import csv
from docx import Document
from pathlib import Path


def get_questions():
    '''
    Reads a csv file with with question statements
    and returns it's content as a dict
    '''
    questions = {}
    with open('questions.csv') as csvfile:
        reader = csv.DictReader(csvfile, delimiter=',')
        for q in reader:
            questions[q['Questão']] = q['Enunciado']

    return questions


def aggregate_responses(reader: csv.DictReader):
    '''
    Aggregates the content of the csv file with survey responses by respondent
    '''
    aggregated = {}
    for response in reader:
        if response['participant_id'] not in aggregated:
            aggregated[response['participant_id']] = []
        aggregated[response['participant_id']].append({
            'Q': response['questionnaireItem_id'],
            'A': response['answer']
        })
    return aggregated


def create_doc(responses: list[dict], questions: dict, name: str):
    '''
    Creates a docx file with questions and answers of a respondent
    '''
    document = Document()
    for r in responses:
        document.add_paragraph().add_run(
            text=f'Q{r['Q']}: {questions[r['Q']]}',
            style='Strong')
        document.add_paragraph(
            text=f'A: {r['A']}'
        )
        document.save(f'output/{int(participant):03}.docx')


if __name__ == "__main__":
    # Create output directory, if not exists
    Path("output").mkdir(exist_ok=True)

    # Get dict with question statements
    questions = get_questions()

    # Create .doc files
    with open('open_questions_responses.csv') as csvfile:
        reader = csv.DictReader(csvfile, delimiter=',')
        aggregated = aggregate_responses(reader)
        for participant in aggregated:
            create_doc(
                responses=aggregated[participant],
                questions=questions,
                name=participant
            )
